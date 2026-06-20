"""Render a Report to a Microsoft Word (.docx) document.

Four modes, chosen automatically from what the template contains:

* **Clone-fill (recommended for your own report format)** — the user marks ONE
  example finding in their template with simple ``[[markers]]`` between
  ``[[finding]]`` and ``[[/finding]]`` (and optionally a summary-table row with
  ``[[findings_row]]``). We deep-copy that example block/row once per finding —
  preserving the company's exact fonts, styles and layout — and substitute the
  markers. Pure python-docx; avoids docxtpl's looping limits on complex
  templates. See the README for the marker reference.

* **Tagged template (docxtpl)** — the ``.docx`` contains Jinja2 placeholders
  (``{{ client }}``, ``{% for f in findings %}`` …); filled with ``docxtpl``.

* **Branding shell (any template)** — a normal template with no markers/tags:
  keep the whole document and *append* the findings as a new section.

* **Default layout** — no template given: a clean report built from scratch.

python-docx / docxtpl are optional; install with
``pip install 'vapt-report-generator[docx]'``.
"""

from __future__ import annotations

from pathlib import Path
from typing import Optional

from ..models import Report

# Severity → RGB accent.
_SEV_RGB = {
    "Critical": (176, 0, 32),
    "High": (211, 84, 0),
    "Medium": (184, 134, 11),
    "Low": (46, 125, 50),
    "Informational": (96, 125, 139),
}

# Table styles we'd like, in order; we fall back to whatever the document has.
_SUMMARY_STYLES = ("Light List Accent 1", "Light List", "Table Grid")
_META_STYLES = ("Light Grid Accent 1", "Light Grid", "Table Grid")


def _context(report: Report) -> dict:
    """Flat, template-friendly context exposed to a user's tagged docx template."""
    return {
        "report": report,
        "client": report.client,
        "title": report.title,
        "assessor": report.assessor,
        "date": report.assessment_date,
        "standard": report.standard,
        "scope": report.scope,
        "risk_rating": report.risk_rating(),
        "counts": report.severity_counts(),
        "findings": report.findings,
        "total": len(report.findings),
    }


def _template_has_placeholders(template_path: str) -> bool:
    """True if the .docx contains any Jinja2 tags ({{ }} or {% %}).

    Run text is rejoined by python-docx, so tags split across runs are detected.
    """
    from docx import Document

    doc = Document(template_path)

    def _iter_text():
        for p in doc.paragraphs:
            yield p.text
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p.text
        for section in doc.sections:
            for hf in (section.header, section.footer):
                for p in hf.paragraphs:
                    yield p.text

    blob = "\n".join(_iter_text())
    return "{{" in blob or "{%" in blob


def _pick_table_style(doc, names):
    available = {s.name for s in doc.styles}
    for n in names:
        if n in available:
            return n
    return None


def _build_findings_body(doc, report) -> None:
    """Append the summary table + per-finding details into ``doc`` using its
    own styles (so it inherits the template's theme/branding)."""
    from docx.shared import RGBColor

    # ---- Summary table ----
    doc.add_heading("Summary of Vulnerability Findings", level=1)
    summary = doc.add_table(rows=1, cols=4)
    style = _pick_table_style(doc, _SUMMARY_STYLES)
    if style:
        summary.style = style
    hdr = summary.rows[0].cells
    for i, h in enumerate(("ID", "Finding", "Severity", "CVSS")):
        hdr[i].text = h
        if hdr[i].paragraphs[0].runs:
            hdr[i].paragraphs[0].runs[0].bold = True
    for f in report.findings:
        c = summary.add_row().cells
        c[0].text = f.finding_id
        c[1].text = f.title
        c[2].text = f.severity.value
        c[3].text = f"{f.cvss_score:.1f}" if f.cvss_score is not None else "—"
        rgb = _SEV_RGB.get(f.severity.value)
        if rgb and c[2].paragraphs[0].runs:
            c[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(*rgb)

    # ---- Per-finding detail ----
    doc.add_heading("Vulnerability Technical Details", level=1)
    for f in report.findings:
        h = doc.add_heading(level=2)
        run = h.add_run(f"{f.finding_id} — {f.title} [{f.severity.value}]")
        rgb = _SEV_RGB.get(f.severity.value)
        if rgb:
            run.font.color.rgb = RGBColor(*rgb)

        bits = []
        if f.cvss_score is not None:
            bits.append(f"CVSS {f.cvss_score:.1f}")
        if f.cvss_vector:
            bits.append(f.cvss_vector)
        if f.cwe:
            bits.append(f.cwe)
        if f.cve:
            bits.append(", ".join(f.cve))
        bits.append(f"Source: {f.source}")
        doc.add_paragraph(" · ".join(bits))

        if f.targets:
            p = doc.add_paragraph()
            p.add_run("Affected: ").bold = True
            p.add_run(", ".join(t.label() for t in f.targets))
        if f.description:
            doc.add_paragraph().add_run("Vulnerability Description:").bold = True
            doc.add_paragraph(f.description)
        if f.evidence:
            doc.add_paragraph().add_run("Proof of Concept / Evidence:").bold = True
            doc.add_paragraph(f.evidence)
        if f.remediation:
            doc.add_paragraph().add_run("Recommendation:").bold = True
            doc.add_paragraph(f.remediation)
        if f.references:
            doc.add_paragraph().add_run("References:").bold = True
            for ref in f.references:
                try:
                    doc.add_paragraph(ref, style="List Bullet")
                except KeyError:
                    doc.add_paragraph(f"• {ref}")


def _render_default(report: Report, output: str) -> str:
    """Build a professional report from scratch (no template needed)."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError(
            "DOCX output requires python-docx. Install with:\n"
            "    pip install 'vapt-report-generator[docx]'"
        ) from exc

    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(report.title)
    run.bold = True
    run.font.size = Pt(24)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"Prepared for: {report.client}").font.size = Pt(14)
    doc.add_paragraph()

    meta = doc.add_table(rows=0, cols=2)
    style = _pick_table_style(doc, _META_STYLES)
    if style:
        meta.style = style
    for k, v in [
        ("Overall Risk", report.risk_rating()),
        ("Assessor", report.assessor),
        ("Date", report.assessment_date),
        ("Standard", report.standard),
        ("Findings", str(len(report.findings))),
        ("Scope", ", ".join(report.scope) if report.scope else "—"),
    ]:
        cells = meta.add_row().cells
        cells[0].text = k
        cells[1].text = v
        if cells[0].paragraphs[0].runs:
            cells[0].paragraphs[0].runs[0].bold = True

    _build_findings_body(doc, report)
    doc.save(output)
    return output


def _render_branding_shell(report: Report, output: str, template_path: str) -> str:
    """Use an untagged company template as a branded shell: keep all of its
    content/styling and append the generated findings as a new section."""
    from docx import Document
    from docx.enum.text import WD_BREAK

    doc = Document(template_path)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)  # start on a fresh page
    _build_findings_body(doc, report)
    doc.save(output)
    return output


def _render_template(report: Report, output: str, template_path: str) -> str:
    """Fill a tagged .docx company template with the report data via docxtpl."""
    try:
        from docxtpl import DocxTemplate
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError(
            "DOCX output requires python-docx + docxtpl. Install with:\n"
            "    pip install 'vapt-report-generator[docx]'"
        ) from exc

    doc = DocxTemplate(template_path)
    doc.render(_context(report))
    doc.save(output)
    return output


# ---------------------------------------------------------------------------
# Clone-fill mode: replicate the user's own example finding block per finding.
#
# The user marks ONE example finding in their template with simple [[markers]]:
#   [[finding]]  ... example finding paragraphs with [[title]], [[cvss]] …  [[/finding]]
# and (optionally) one summary-table row containing [[findings_row]].
# We deep-copy that block/row once per finding — so the company's exact fonts,
# styles and layout are preserved — and substitute the markers. This is the
# python-docx way; it avoids docxtpl's loop limitations on complex templates.
# ---------------------------------------------------------------------------
import re as _re

_MARK = _re.compile(r"\[\[([a-z_]+)\]\]")
_BLOCK_START = "[[finding]]"
_BLOCK_END = "[[/finding]]"
_ROW_MARK = "[[findings_row]]"


def _finding_markers(f, sn: int) -> dict:
    cvss = f"{f.cvss_score:.1f}" if f.cvss_score is not None else "N/A"
    return {
        "sn": str(sn),
        "id": f.finding_id,
        "title": f.title or "N/A",
        "severity": f.severity.value,
        "cvss": cvss,
        "cvss_vector": f.cvss_vector or "N/A",
        "cwe": f.cwe or "N/A",
        "cve": ", ".join(f.cve) if f.cve else "N/A",
        "targets": ", ".join(t.label() for t in f.targets) or "N/A",
        "description": (f.description or "N/A").strip(),
        "evidence": (f.evidence or "N/A").replace("\r", " ").strip(),
        "remediation": (f.remediation or "N/A").strip(),
        "references": ", ".join(f.references) if f.references else "N/A",
    }


def _doc_markers(report: Report) -> dict:
    c = report.severity_counts()
    return {
        "client": report.client,
        "report_title": report.title,
        "assessor": report.assessor,
        "date": report.assessment_date,
        "standard": report.standard,
        "risk_rating": report.risk_rating(),
        "total_findings": str(len(report.findings)),
        "count_critical": str(c["Critical"]),
        "count_high": str(c["High"]),
        "count_medium": str(c["Medium"]),
        "count_low": str(c["Low"]),
    }


def _fill_markers(paragraph, mapping: dict) -> None:
    """Replace [[markers]] in a paragraph, preserving runs/breaks where the
    marker sits inside a single run; falling back to a paragraph-level set
    (keeps the paragraph style + first run font) for run-split markers."""
    # First pass: run-level (keeps intra-paragraph formatting and line breaks).
    for run in paragraph.runs:
        if "[[" in run.text:
            run.text = _MARK.sub(lambda m: mapping.get(m.group(1), m.group(0)), run.text)
    # Second pass: any marker split across runs -> rebuild paragraph text.
    joined = "".join(r.text for r in paragraph.runs)
    if "[[" in joined and _MARK.search(joined):
        new = _MARK.sub(lambda m: mapping.get(m.group(1), m.group(0)), joined)
        if new != joined and paragraph.runs:
            paragraph.runs[0].text = new
            for r in paragraph.runs[1:]:
                r.text = ""


def _template_has_clone_markers(template_path: str) -> bool:
    from docx import Document

    doc = Document(template_path)
    for p in doc.paragraphs:
        if _BLOCK_START in p.text or _ROW_MARK in p.text:
            return True
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                if _ROW_MARK in cell.text or _BLOCK_START in cell.text:
                    return True
    return False


def _render_clone_fill(report: Report, output: str, template_path: str) -> str:
    """Clone the marked example finding block / summary row once per finding."""
    import copy

    from docx import Document
    from docx.text.paragraph import Paragraph

    doc = Document(template_path)
    body = doc._body
    findings = report.findings

    # --- 1) per-finding detail block between [[finding]] and [[/finding]] ---
    paras = doc.paragraphs
    starts = [p for p in paras if p.text.strip() == _BLOCK_START]
    ends = [p for p in paras if p.text.strip() == _BLOCK_END]
    if starts and ends:
        start_el, end_el = starts[0]._p, ends[0]._p
        # paragraphs strictly between the two markers form the template block
        block_els, cur = [], start_el.getnext()
        while cur is not None and cur is not end_el:
            block_els.append(cur)
            cur = cur.getnext()
        template_block = [copy.deepcopy(el) for el in block_els]
        for idx, f in enumerate(findings, 1):
            mapping = _finding_markers(f, idx)
            for el in (copy.deepcopy(e) for e in template_block):
                end_el.addprevious(el)
                _fill_markers(Paragraph(el, body), mapping)
        for el in (*block_els, start_el, end_el):
            el.getparent().remove(el)

    # --- 2) summary-table row containing [[findings_row]] ---
    for table in doc.tables:
        marker_row = None
        for row in table.rows:
            if any(_ROW_MARK in c.text for c in row.cells):
                marker_row = row
                break
        if marker_row is None:
            continue
        tmpl_tr = copy.deepcopy(marker_row._tr)
        for idx, f in enumerate(findings, 1):
            new_tr = copy.deepcopy(tmpl_tr)
            marker_row._tr.addprevious(new_tr)
            from docx.table import _Row

            new_row = _Row(new_tr, table)
            mapping = _finding_markers(f, idx)
            for cell in new_row.cells:
                for p in cell.paragraphs:
                    # drop the [[findings_row]] sentinel, then fill fields
                    for run in p.runs:
                        run.text = run.text.replace(_ROW_MARK, "")
                    _fill_markers(p, mapping)
        marker_row._tr.getparent().remove(marker_row._tr)

    # --- 3) document-level single markers (client/date/counts…) ---
    docmap = _doc_markers(report)
    for p in doc.paragraphs:
        if "[[" in p.text:
            _fill_markers(p, docmap)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if "[[" in p.text:
                        _fill_markers(p, docmap)

    doc.save(output)
    return output


# ---------------------------------------------------------------------------
# Auto-fill mode: detect the template's OWN example finding block (no markers,
# no tags) and replicate it per finding, in place, in the exact formatting.
# This is what makes a plain company report template "just work": the example
# finding under a "Technical Details"/"Findings" heading is recognised by its
# field labels (Description / CVSS / CWE / Severity / Affected / Remediation /
# References), cloned once per finding and filled — keeping everything inside
# the template instead of appending below it.
# ---------------------------------------------------------------------------
_HEADING1 = ("Heading 1", "Title", "Style1")
_HEADING_SUB = ("Heading 2", "Heading 3", "Heading 4", "Style2", "Style3")
_SECTION_KEYS = ("technical detail", "detailed finding", "vulnerability detail",
                 "findings detail", "finding detail")


def _is_h1(p) -> bool:
    return p.style.name.startswith(("Heading 1", "Style1")) or p.style.name == "Title"


def _set_keep(p, text: str) -> None:
    """Value-only paragraph: set text, keep the first run's font + paragraph style."""
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def _set_inline(p, value: str) -> None:
    """'Label: value' paragraph: keep the bold label + colon, replace the value."""
    runs = p.runs
    ci = next((i for i, r in enumerate(runs) if ":" in r.text), None)
    if ci is None:
        _set_keep(p, value)
        return
    if ci + 1 < len(runs):
        runs[ci + 1].text = " " + value
        for r in runs[ci + 2:]:
            r.text = ""
    else:
        p.add_run(" " + value)


def _fill_affected(p, targets: str) -> None:
    from docx.oxml.ns import qn

    for hl in p._p.findall(qn("w:hyperlink")):
        p._p.remove(hl)
    runs = p.runs
    nl = next((i for i, r in enumerate(runs) if "\n" in r.text), len(runs))
    ci = next((i for i, r in enumerate(runs) if ":" in r.text), 0)
    placed = False
    for i in range(ci + 1, nl):
        if not placed:
            runs[i].text = " " + targets
            placed = True
        else:
            runs[i].text = ""
    if not placed:
        if ci + 1 <= len(runs) - 1:
            runs[ci + 1].text = " " + targets
        else:
            p.add_run(" " + targets)


def _autodetect_block(doc):
    """Return (start, end) paragraph indices of the example finding block, or None."""
    ps = doc.paragraphs
    tech = next((i for i, p in enumerate(ps)
                 if _is_h1(p) and any(k in p.text.lower() for k in _SECTION_KEYS)), None)
    if tech is None:
        return None
    start = next((i for i in range(tech + 1, len(ps))
                  if ps[i].text.strip() and ps[i].style.name in _HEADING_SUB), None)
    if start is None:
        return None
    end = len(ps) - 1
    for i in range(start + 1, len(ps)):
        if _is_h1(ps[i]) or ps[i].text.strip().lower().startswith("appendix"):
            end = i - 1
            break
    while end > start and not ps[end].text.strip():
        end -= 1
    # require at least a couple of recognisable field labels to be confident
    labels = " ".join(ps[i].text.lower() for i in range(start, end + 1))
    if sum(k in labels for k in ("description", "severity", "remediation", "cvss", "cwe")) < 2:
        return None
    return start, end


def _fill_finding_block(block, f) -> None:
    """Fill one cloned finding block (list of Paragraphs) with finding ``f``."""
    _set_keep(block[0], f.title or "N/A")
    i = 1
    while i < len(block):
        p = block[i]
        low = p.text.strip().lower()
        if low.endswith("description:"):
            _set_keep(block[i + 1], (f.description or "N/A").strip())
            i += 2
            continue
        if low.startswith("cvss") and low.endswith(":"):
            cv = f"{f.cvss_score:.1f}" if f.cvss_score is not None else "N/A"
            if f.cvss_vector:
                cv += f"  ({f.cvss_vector})"
            _set_keep(block[i + 1], cv)
            i += 2
            continue
        if ("remediation" in low or "recommendation" in low) and low.endswith(":"):
            _set_keep(block[i + 1], (f.remediation or "N/A").strip())
            j = i + 2
            while j < len(block) and not block[j].text.strip().lower().startswith("reference"):
                _set_keep(block[j], "")
                j += 1
            i = j
            continue
        if low.startswith("cwe"):
            _set_inline(p, f.cwe or "N/A")
        elif low.startswith("severity"):
            _set_inline(p, f.severity.value)
        elif low.startswith("affected"):
            _fill_affected(p, ", ".join(t.label() for t in f.targets) or "N/A")
        elif low.startswith("reference"):
            _set_inline(p, ", ".join(f.references) if f.references else "N/A")
        elif p.style.name == "List Paragraph":
            _set_keep(p, (f.evidence or "N/A").replace("\r", " ").strip())
            j = i + 1
            while j < len(block) and block[j].style.name == "List Paragraph":
                block[j]._p.getparent().remove(block[j]._p)
                j += 1
            i = j
            continue
        i += 1


def _render_autofill(report: Report, output: str, template_path: str, block) -> str:
    """Clone+fill the detected example finding block per finding, and refresh the
    summary table, keeping everything inside the template."""
    import copy

    from docx import Document
    from docx.text.paragraph import Paragraph

    doc = Document(template_path)
    body = doc._body
    findings = report.findings

    ps = doc.paragraphs
    start, end = block
    block_els = [ps[i]._p for i in range(start, end + 1)]
    template_block = [copy.deepcopy(e) for e in block_els]

    _fill_finding_block([Paragraph(e, body) for e in block_els], findings[0])
    prev = block_els[-1]
    for f in findings[1:]:
        new = [copy.deepcopy(e) for e in template_block]
        ref = prev
        for e in new:
            ref.addnext(e)
            ref = e
        _fill_finding_block([Paragraph(e, body) for e in new], f)
        prev = new[-1]

    # summary table: the widest "findings" table (most rows / 4 cols) gets refreshed
    cand = [t for t in doc.tables if len(t.columns) == 4 and t.rows]
    if cand:
        tbl = max(cand, key=lambda t: len(t.rows))
        tmpl_tr = copy.deepcopy(tbl.rows[0]._tr)
        for row in list(tbl.rows):
            row._tr.getparent().remove(row._tr)
        for sn, f in enumerate(findings, 1):
            tbl._tbl.append(copy.deepcopy(tmpl_tr))
            c = tbl.rows[-1].cells
            _set_keep(c[0].paragraphs[0], str(sn))
            _set_keep(c[1].paragraphs[0], f.title)
            _set_keep(c[2].paragraphs[0], f.severity.value)
            _set_keep(c[3].paragraphs[0], (f.description or "N/A")[:200])

    doc.save(output)
    return output


def render(report: Report, output: str, template_path: Optional[str] = None) -> str:
    """Render ``report`` to a .docx file.

    * no template               → default professional layout
    * template with [[markers]]  → clone-fill the marked example block
    * template with {{ }} tags   → precise fill via docxtpl
    * template with a detectable example finding block → auto-fill it in place
    * otherwise                  → branding shell: keep it, append the findings
    """
    if not template_path:
        return _render_default(report, output)

    if not Path(template_path).exists():
        raise FileNotFoundError(f"Template not found: {template_path}")
    if not template_path.lower().endswith(".docx"):
        raise ValueError(
            f"For DOCX output a custom template must be a .docx file, not "
            f"'{Path(template_path).name}'."
        )

    if _template_has_clone_markers(template_path):
        return _render_clone_fill(report, output, template_path)
    if _template_has_placeholders(template_path):
        return _render_template(report, output, template_path)

    if report.findings:
        from docx import Document

        block = _autodetect_block(Document(template_path))
        if block is not None:
            return _render_autofill(report, output, template_path, block)

    return _render_branding_shell(report, output, template_path)
