"""Unit tests for vapt-report-generator."""

from pathlib import Path

import pytest

from vaptreport import scoring
from vaptreport.models import Finding, Report, Severity, Target
from vaptreport.parsers import detect_and_parse

EXAMPLES = Path(__file__).resolve().parent.parent / "examples"


# ── CVSS scoring ───────────────────────────────────────────────────────────────
@pytest.mark.parametrize(
    "vector,expected",
    [
        ("CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:H/A:H", 9.8),   # EternalBlue
        ("CVSS:3.1/AV:N/AC:H/PR:N/UI:N/S:U/C:H/I:N/A:N", 5.9),
        ("CVSS:3.1/AV:L/AC:L/PR:L/UI:N/S:U/C:N/I:N/A:N", 0.0),
        ("CVSS:3.1/AV:N/AC:L/PR:N/UI:R/S:C/C:H/I:L/A:N", 8.2),
    ],
)
def test_base_score(vector, expected):
    assert scoring.base_score(vector) == expected


def test_severity_from_score():
    assert scoring.severity_from_score(9.8) == "Critical"
    assert scoring.severity_from_score(7.0) == "High"
    assert scoring.severity_from_score(4.0) == "Medium"
    assert scoring.severity_from_score(0.1) == "Low"
    assert scoring.severity_from_score(0.0) == "Informational"


def test_invalid_vector_raises():
    with pytest.raises(ValueError):
        scoring.base_score("CVSS:3.1/AV:N/AC:L")  # missing metrics


# ── Models ─────────────────────────────────────────────────────────────────────
def test_finding_derives_score_from_vector():
    f = Finding(title="x", cvss_vector="CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:H/A:H")
    assert f.cvss_score == 9.8
    assert f.severity == Severity.CRITICAL


def test_report_sort_and_ids():
    report = Report(findings=[
        Finding(title="low", severity=Severity.LOW),
        Finding(title="crit", severity=Severity.CRITICAL),
        Finding(title="med", severity=Severity.MEDIUM),
    ])
    report.finalize()
    assert [f.title for f in report.findings] == ["crit", "med", "low"]
    assert report.findings[0].finding_id == "VR-001"
    assert report.risk_rating() == "Critical"


def test_severity_coerce():
    assert Severity.coerce("HIGH") == Severity.HIGH
    assert Severity.coerce("moderate") == Severity.MEDIUM
    assert Severity.coerce(None) == Severity.INFO


def test_target_label():
    assert Target("h", 443, "https").label() == "h:443 (https)"
    assert Target("h").label() == "h"


# ── Parsers ────────────────────────────────────────────────────────────────────
def test_parse_findings_json():
    findings = detect_and_parse(str(EXAMPLES / "sample_findings.json"))
    assert len(findings) == 5
    sqli = next(f for f in findings if "SQL Injection" in f.title)
    assert sqli.severity == Severity.CRITICAL
    assert sqli.cvss_score == 9.8


def test_no_acme_in_examples():
    text = (EXAMPLES / "sample_findings.json").read_text().lower()
    assert "acme" not in text


def test_parse_nmap():
    findings = detect_and_parse(str(EXAMPLES / "sample_nmap.xml"))
    # POODLE vuln script + open-service inventory
    assert any("ssl-poodle" in f.title for f in findings)
    assert any("Open Network Services" in f.title for f in findings)


# Synthetic Acunetix Developer Report text (mimics the real PDF layout,
# including the typographic ligatures Acunetix uses: "Classiﬁcation",
# "Aﬀected"). Lets us test the parser without shipping a real client PDF.
_ACUNETIX_TEXT = """Developer Report
Acunetix Security Audit
Scan of test.example.com
Host test.example.com
Start url https://test.example.com/
2
Alerts summary
Conﬁguration ﬁle disclosure
Classiﬁcation
CVSS3
CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:C/C:L/I:N/A:NBase Score: 5.8Attack Vector: Network
CWE CWE-538
Aﬀected items Variation
/WEB-INF/web.xml 1
HTTP Strict Transport Security (HSTS) Policy Not Enabled
Classiﬁcation
CWE CWE-16
Aﬀected items Variation
Web Server 1
3
Alerts details
Conﬁguration ﬁle disclosure
Severity High
Reported by module /Scripts/test.script
Description
A backup conﬁguration ﬁle was found.
Impact
Discloses sensitive information.
Recommendation
Remove this ﬁle from the web server.
References
OWASP (https://owasp.org/test)
Aﬀected items
/WEB-INF/web.xml
Details
Pattern found: <web-app>OWIT</web-app>
Request headers
GET /WEB-INF/web.xml HTTP/1.1
4
HTTP Strict Transport Security (HSTS) Policy Not Enabled
Severity Medium
Reported by module /httpdata/HSTS.js
Description
The Strict-Transport-Security header is missing.
Impact
MitM attacks possible.
Recommendation
Implement HSTS.
References
hstspreload (https://hstspreload.org/)
Aﬀected items
Web Server
Details
URLs where HSTS is not enabled: https://test.example.com/a/
Request headers
GET /a/ HTTP/1.1
5
Scanned items (coverage report)
https://test.example.com/
"""


def test_parse_acunetix_text():
    from vaptreport.parsers import acunetix_pdf

    findings = acunetix_pdf.parse_text(_ACUNETIX_TEXT)
    assert len(findings) == 2

    cfg = findings[0]
    assert cfg.title == "Configuration file disclosure"   # ligature normalised
    assert cfg.severity == Severity.HIGH
    assert cfg.cvss_score == 5.8                            # parsed from summary
    assert cfg.cwe == "CWE-538"
    assert "backup configuration file" in cfg.description
    assert "Remove this file" in cfg.remediation
    assert cfg.references == ["https://owasp.org/test"]
    assert "Pattern found" in cfg.evidence
    assert cfg.targets[0].host == "test.example.com"
    assert cfg.source == "acunetix"

    hsts = findings[1]
    assert hsts.severity == Severity.MEDIUM
    assert hsts.cwe == "CWE-16"
    assert hsts.cvss_score is None                          # no CVSS in summary


def test_parse_nessus_merges_hosts():
    findings = detect_and_parse(str(EXAMPLES / "sample_nessus.nessus"))
    eternalblue = next(f for f in findings if "EternalBlue" in f.title)
    # Same plugin on two hosts -> one finding, two targets.
    assert len(eternalblue.targets) == 2
    assert eternalblue.severity == Severity.CRITICAL


def test_parse_nuclei_jsonl():
    findings = detect_and_parse(str(EXAMPLES / "sample_nuclei.jsonl"))
    assert len(findings) == 4                      # git-config merged across 2 hosts
    log4shell = next(f for f in findings if "Log4" in f.title)
    assert log4shell.severity == Severity.CRITICAL
    assert log4shell.cvss_score == 10.0
    assert log4shell.cve == ["CVE-2021-44228"]
    gitcfg = next(f for f in findings if "Git Config" in f.title)
    assert len(gitcfg.targets) == 2


def test_parse_burp_xml_merges_paths():
    findings = detect_and_parse(str(EXAMPLES / "sample_burp.xml"))
    assert len(findings) == 4                      # XSS merged across 2 paths
    xss = next(f for f in findings if "scripting" in f.title.lower())
    assert xss.severity == Severity.HIGH
    assert xss.cwe == "CWE-79"
    assert len(xss.targets) == 2
    info = next(f for f in findings if "banner" in f.title.lower())
    assert info.severity == Severity.INFO          # Burp "Information" -> Informational


def test_parse_zap_json():
    findings = detect_and_parse(str(EXAMPLES / "sample_zap.json"))
    assert len(findings) == 4
    sqli = next(f for f in findings if "SQL" in f.title)
    assert sqli.severity == Severity.HIGH          # riskcode 3
    assert sqli.cwe == "CWE-89"
    csp = next(f for f in findings if "CSP" in f.title)
    assert len(csp.targets) == 2                    # two instances
    assert "parameterized queries" in sqli.remediation.lower()


def test_parse_zap_xml(tmp_path):
    from vaptreport.parsers import detect_and_parse as dp
    xml = """<?xml version="1.0"?>
<OWASPZAPReport version="2.15.0">
  <site name="https://app.example.com" host="app.example.com" port="443" ssl="true">
    <alerts>
      <alertitem>
        <alert>Path Traversal</alert>
        <riskcode>3</riskcode>
        <desc>&lt;p&gt;Path traversal may be possible.&lt;/p&gt;</desc>
        <solution>&lt;p&gt;Validate and canonicalize file paths.&lt;/p&gt;</solution>
        <reference>&lt;p&gt;https://owasp.org/path-traversal&lt;/p&gt;</reference>
        <cweid>22</cweid>
        <instances><instance><uri>https://app.example.com/download?file=../../etc/passwd</uri><method>GET</method><param>file</param></instance></instances>
      </alertitem>
    </alerts>
  </site>
</OWASPZAPReport>"""
    f = tmp_path / "zap.xml"
    f.write_text(xml)
    findings = dp(str(f))
    assert len(findings) == 1
    assert findings[0].title == "Path Traversal"
    assert findings[0].severity == Severity.HIGH
    assert findings[0].cwe == "CWE-22"
    assert findings[0].source == "zap"


def test_json_dialect_sniffing():
    """All three JSON dialects route to the correct parser."""
    from vaptreport.parsers import _sniff_json
    assert _sniff_json(str(EXAMPLES / "sample_zap.json")) == "zap"
    assert _sniff_json(str(EXAMPLES / "sample_findings.json")) == "findings"


def test_xml_dialect_sniffing():
    from vaptreport.parsers import _sniff_xml
    assert _sniff_xml(str(EXAMPLES / "sample_nmap.xml")) == "nmap"
    assert _sniff_xml(str(EXAMPLES / "sample_burp.xml")) == "burp"


# ── Reporters ──────────────────────────────────────────────────────────────────
def test_render_html(tmp_path):
    from vaptreport import reporters

    findings = detect_and_parse(str(EXAMPLES / "sample_findings.json"))
    report = Report(findings=findings).finalize()
    out = reporters.render(report, "html", str(tmp_path / "r.html"))
    html = Path(out).read_text()
    assert "SQL Injection" in html
    assert "VR-001" in html


def test_html_escapes_xss_payload(tmp_path):
    """An XSS PoC in an evidence field must be escaped, not executable."""
    from vaptreport import reporters

    findings = detect_and_parse(str(EXAMPLES / "sample_findings.json"))
    report = Report(findings=findings).finalize()
    html = Path(reporters.render(report, "html", str(tmp_path / "r.html"))).read_text()
    # The raw, executable payload must NOT appear...
    assert "<img src=x onerror=alert" not in html
    # ...but its escaped, inert form must.
    assert "&lt;img src=x onerror=alert" in html


def test_custom_template(tmp_path):
    from vaptreport import reporters

    findings = detect_and_parse(str(EXAMPLES / "sample_findings.json"))
    report = Report(findings=findings).finalize()
    out = reporters.render(
        report, "html", str(tmp_path / "r.html"),
        template=str(EXAMPLES / "custom_template.html.j2"),
    )
    html = Path(out).read_text()
    assert "SENTINEL SECURITY" in html  # branding from the custom template
    assert "VR-001" in html


def test_custom_template_not_supported_for_xlsx(tmp_path):
    from vaptreport import reporters

    report = Report(findings=detect_and_parse(str(EXAMPLES / "sample_findings.json")))
    with pytest.raises(ValueError):
        reporters.render(report, "xlsx", str(tmp_path / "r.xlsx"), template="x.j2")


def test_pdf_passed_as_template_gives_clear_error(tmp_path):
    from vaptreport import reporters

    # A user pointing --template at their company report PDF must get a helpful
    # message, not a raw UTF-8 decode crash.
    pdf = tmp_path / "company template.pdf"
    pdf.write_bytes(b"%PDF-1.7\n\xb5\xb5\xb5 binary stream \x00\x01")
    report = Report(findings=detect_and_parse(str(EXAMPLES / "sample_findings.json")))
    with pytest.raises(ValueError, match="PDF"):
        reporters.render(report, "html", str(tmp_path / "r.html"), template=str(pdf))


def test_binary_template_gives_clear_error(tmp_path):
    from vaptreport import reporters

    blob = tmp_path / "logo.png"
    blob.write_bytes(b"\x89PNG\r\n\x1a\n\xb5\xff\x00not text")
    report = Report(findings=detect_and_parse(str(EXAMPLES / "sample_findings.json")))
    with pytest.raises(ValueError, match="UTF-8 text"):
        reporters.render(report, "html", str(tmp_path / "r.html"), template=str(blob))


def test_render_pdf(tmp_path):
    pytest.importorskip("weasyprint")
    from vaptreport import reporters

    findings = detect_and_parse(str(EXAMPLES / "sample_findings.json"))
    report = Report(findings=findings).finalize()
    out = reporters.render(report, "pdf", str(tmp_path / "r.pdf"))
    data = Path(out).read_bytes()
    assert data[:4] == b"%PDF"  # valid PDF magic bytes


def test_render_xlsx(tmp_path):
    from vaptreport import reporters

    findings = detect_and_parse(str(EXAMPLES / "sample_findings.json"))
    report = Report(findings=findings).finalize()
    out = reporters.render(report, "xlsx", str(tmp_path / "r.xlsx"))
    assert Path(out).exists() and Path(out).stat().st_size > 0


def _docx_text(path) -> str:
    from docx import Document

    return "\n".join(p.text for p in Document(str(path)).paragraphs)


def test_render_docx_default(tmp_path):
    pytest.importorskip("docx")
    from vaptreport import reporters

    findings = detect_and_parse(str(EXAMPLES / "sample_findings.json"))
    report = Report(findings=findings).finalize()
    out = reporters.render(report, "docx", str(tmp_path / "r.docx"))
    text = _docx_text(out)
    assert "Summary of Vulnerability Findings" in text
    assert "SQL Injection" in text  # a finding from the sample


def test_render_docx_with_company_template(tmp_path):
    pytest.importorskip("docxtpl")
    from vaptreport import reporters

    findings = detect_and_parse(str(EXAMPLES / "sample_findings.json"))
    report = Report(findings=findings, client="OWIT Global").finalize()
    out = reporters.render(report, "docx", str(tmp_path / "r.docx"),
                           template=str(EXAMPLES / "company_template.docx"))
    text = _docx_text(out)
    assert "OWIT Global" in text          # field filled from the report
    assert "{{" not in text and "{%" not in text  # no unrendered placeholders


def test_docx_untagged_template_becomes_branding_shell(tmp_path):
    pytest.importorskip("docx")
    from docx import Document
    from vaptreport import reporters

    # A company template that is just a formatted doc with NO {{ }} / {% %} tags.
    plain = tmp_path / "template.docx"
    d = Document()
    d.add_heading("OWIT Global — Penetration Test Report", level=0)
    d.add_paragraph("This is our standard report layout and branding.")
    d.save(str(plain))

    report = Report(findings=detect_and_parse(str(EXAMPLES / "sample_findings.json"))).finalize()
    out = reporters.render(report, "docx", str(tmp_path / "out.docx"), template=str(plain))
    text = _docx_text(out)
    # The template's own content is preserved…
    assert "OWIT Global — Penetration Test Report" in text
    assert "standard report layout" in text
    # …and the findings are appended in a new section.
    assert "Summary of Vulnerability Findings" in text
    assert "SQL Injection" in text


def test_docx_template_must_be_docx(tmp_path):
    pytest.importorskip("docx")
    from vaptreport import reporters

    report = Report(findings=detect_and_parse(str(EXAMPLES / "sample_findings.json")))
    with pytest.raises(ValueError, match="must be a .docx"):
        reporters.render(report, "docx", str(tmp_path / "r.docx"),
                         template=str(EXAMPLES / "custom_template.html.j2"))


def test_pdf_template_reuses_cover_page(tmp_path):
    pytest.importorskip("weasyprint")
    pytest.importorskip("pypdf")
    from pypdf import PdfReader
    from weasyprint import HTML
    from vaptreport import reporters

    cover = tmp_path / "company_cover.pdf"
    HTML(string="<h1>ACME COVER PAGE</h1>").write_pdf(str(cover))
    report = Report(findings=detect_and_parse(str(EXAMPLES / "sample_findings.json"))).finalize()
    out = reporters.render(report, "pdf", str(tmp_path / "r.pdf"), template=str(cover))
    reader = PdfReader(out)
    assert "ACME COVER PAGE" in reader.pages[0].extract_text()  # cover preserved
    assert len(reader.pages) > 1  # findings appended after the cover
